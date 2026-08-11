from dashborad_styles import load_css

from pathlib import Path
import sys
import markdown
import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

import streamlit.components.v1 as components
import pandas as pd
import plotly.express as px
import streamlit as st
import requests
import asyncio
from mcp_configuration.agent import InsuranceAgent

def add_formatted_text(paragraph, text):

    # Split around bold Markdown
    parts = re.split(
        r"(\*\*.*?\*\*)",
        text
    )

    for part in parts:

        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):

            run = paragraph.add_run(
                part[2:-2]
            )

            run.bold = True

        else:

            paragraph.add_run(part)

def markdown_to_word(markdown_text):
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10)

    lines = markdown_text.splitlines()

    i = 0

    while i < len(lines):

        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # --------------------------------
        # Headings
        # --------------------------------

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)

        if heading_match:
            level = len(heading_match.group(1))
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", heading_match.group(2))

            paragraph = doc.add_heading(
                text,
                level=min(level, 6)
            )

            i += 1
            continue

        # --------------------------------
        # Markdown table
        # --------------------------------

        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and "|" in lines[i + 1]
            and re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1])
        ):

            headers = [
                cell.strip()
                for cell in line.strip("|").split("|")
            ]

            i += 2  # Skip header + separator

            rows = []

            while i < len(lines) and lines[i].strip().startswith("|"):

                row = [
                    cell.strip()
                    for cell in lines[i].strip("|").split("|")
                ]

                rows.append(row)
                i += 1

            table = doc.add_table(
                rows=1,
                cols=len(headers)
            )

            table.style = "Table Grid"

            # Header
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = re.sub(
                    r"\*\*(.*?)\*\*",
                    r"\1",
                    header
                )

                for run in cell.paragraphs[0].runs:
                    run.bold = True

            # Rows
            for row in rows:

                cells = table.add_row().cells

                for j, value in enumerate(row):

                    if j < len(cells):
                        cells[j].text = re.sub(
                            r"\*\*(.*?)\*\*",
                            r"\1",
                            value
                        )

            doc.add_paragraph()

            continue

        # --------------------------------
        # Bullet points
        # --------------------------------

        if re.match(r"^[-*]\s+", line):

            text = re.sub(
                r"^[-*]\s+",
                "",
                line
            )

            paragraph = doc.add_paragraph(
                style="List Bullet"
            )

            add_formatted_text(paragraph, text)

            i += 1
            continue

        # --------------------------------
        # Numbered list
        # --------------------------------

        if re.match(r"^\d+\.\s+", line):

            text = re.sub(
                r"^\d+\.\s+",
                "",
                line
            )

            paragraph = doc.add_paragraph(
                style="List Number"
            )

            add_formatted_text(paragraph, text)

            i += 1
            continue

        # --------------------------------
        # Horizontal rule
        # --------------------------------

        if re.match(r"^-{3,}$", line):

            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)

            i += 1
            continue

        # --------------------------------
        # Normal paragraph
        # --------------------------------

        paragraph = doc.add_paragraph()

        add_formatted_text(
            paragraph,
            line
        )

        i += 1

    # Save to memory
    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return output



def clean_markdown_tables(text: str) -> str:
    lines = text.splitlines()
    cleaned = []

    in_table = False

    for line in lines:
        stripped = line.rstrip()

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):

            # Ensure a blank line before the table
            if not in_table and cleaned and cleaned[-1] != "":
                cleaned.append("")

            in_table = True
            cleaned.append(stripped)
            continue

        # Remove blank lines inside a table
        if in_table and stripped == "":
            continue

        # Leaving a table
        if in_table:
            in_table = False
            cleaned.append("")

        cleaned.append(line)

    return "\n".join(cleaned)

# ==========================================================
# Project Imports
# ==========================================================

if "agent" not in st.session_state:
    st.session_state.agent = InsuranceAgent()

agent = st.session_state.agent

DASHBOARD_DIR = Path(__file__).resolve().parent.parent
ROOT_DIR = DASHBOARD_DIR.parent

sys.path.insert(0, str(DASHBOARD_DIR))
sys.path.insert(0, str(ROOT_DIR))

from dashboard.services.llm.analyze_channel import analyze_channel

load_css()

# ==========================================================
# Load Dataset
# ==========================================================

DATASET = (
    ROOT_DIR
    / "Synthetic_Generator"
    / "data"
    / "April_month_renewals_sheet.csv"
)


@st.cache_data
def load_data():

    return pd.read_csv(DATASET)


df = load_data()



# ==========================================================
# Page
# ==========================================================

st.title("📊 Channel Analysis Dashboard")

st.caption(
    "Executive overview of channel performance followed by individual channel analytics."
)


# ==========================================================
# Executive Overview
# ==========================================================

st.header("Executive Overview")

kpi1, kpi2, kpi3 = st.columns(3)

kpi1.metric(

    "Total Policies",

    f"{len(df):,}"

)

kpi2.metric(

    "Total IMDs",

    # df["IMD_Code"].nunique()
    df["New_IMD_Code"].nunique()

)

kpi3.metric(

    "Overall Renewal Rate",

    # f"{df['Renewed'].mean()*100:.1f}%"
    f"{(df['STATUS'] == 'Renewed').mean()*100:.1f}%"

)

# kpi4.metric(

#     "Average Premium",

#     f"₹ {df['Premium'].mean()/1000:,.1f}K"

# )

st.divider()

channel_summary = (

    df

    .groupby("IMD_Channel")

    .agg(

        Policies=("Policy_Number", "count"),

        Renewal_Rate=("STATUS", lambda s: (s == "Renewed").mean() * 100),

        # Average_Premium=("Premium", "mean"),

        # Average_Claims=("Claim_Count", "mean")

    )

    .reset_index()

)

# channel_summary["Renewal_Rate"] *= 100

left, right = st.columns([2, 1])

with left:

    st.subheader("Renewal Rate by Channel")

    fig = px.bar(

        channel_summary.sort_values(
            "Renewal_Rate"
        ),

        x="Renewal_Rate",

        y="IMD_Channel",

        orientation="h",

        text="Renewal_Rate"

    )

    fig.update_traces(

        texttemplate="%{text:.1f}%"

    )

    fig.update_layout(

        xaxis_title="Renewal Rate (%)",

        yaxis_title="",

        height=350

    )

    st.plotly_chart(

        fig,

        use_container_width=True

    )


with right:

    st.subheader("Business Distribution")

    fig = px.pie(

        channel_summary,

        values="Policies",

        names="IMD_Channel",

        hole=0.5

    )

    st.plotly_chart(

        fig,

        use_container_width=True

    )

st.subheader("Channel Comparison")

display = channel_summary.copy()

display["Renewal_Rate"] = (
    display["Renewal_Rate"]
    .round(1)
)

# display["Average_Premium"] = (
#     display["Average_Premium"]
#     .round(0)
# )

# display["Average_Claims"] = (
#     display["Average_Claims"]
#     .round(2)
# )

st.dataframe(

    display,

    use_container_width=True,

    hide_index=True

)

st.divider()

df["Renewal_flag"] = (
    df["Renewal_flag"].str.extract(r"(\d+)")[0].astype(int)
)

# ==========================================================
# Individual Channel Analysis
# ==========================================================

for channel in sorted(df["IMD_Channel"].unique()):

    st.header(f"📌 {channel} Performance")

    channel_df = df[
        df["IMD_Channel"] == channel
    ].copy()

    # ======================================================
    # Channel KPIs
    # ======================================================

    kpi1, kpi2, kpi3  = st.columns(3)

    kpi1.metric(

        "Policies",

        len(channel_df)

    )

    kpi2.metric(

        "IMDs",

        channel_df["New_IMD_Code"].nunique()

    )

    kpi3.metric(

        "Renewal Rate",

        f"{(df['STATUS'] == 'Renewed').mean()*100:.1f}%"

    )

    # kpi4.metric(

    #     "Average Premium",

    #     f"₹ {channel_df['Premium'].mean()/1000:,.1f}K"

    # )

    st.divider()


    # ======================================================
    # IMD Summary
    # ======================================================

    imd_summary = (

        channel_df

        .groupby("New_IMD_Code")

        .agg(

            Portfolio_Size=("Policy_Number", "count"),

            Renewal_Rate = ("STATUS", lambda s: (s == "Renewed").mean() * 100),

            # Average_Premium=("Premium", "mean"),

            # Total_Claims=("Claim_Count", "sum"),

            # Average_NCB=("NCB", "mean"),

            Average_Vehicle_Age = ("TBR_Veh_Age", "mean"),

            Average_Tenure=("Renewal_flag", "mean")

        )

        .reset_index()

    )

    # imd_summary["Renewal_Rate"] *= 100

    # ======================================================
    # Top & Bottom IMDs
    # ======================================================

    imd_summary["New_IMD_Code"] = imd_summary["New_IMD_Code"].astype(str)

    left, right = st.columns(2)

    with left:

        st.subheader("🟢 Top Performing IMDs")

        top = (

            imd_summary

            .sort_values(

                "Renewal_Rate",

                ascending=False

            )

            .head(10)

        )

        fig = px.bar(

            top.sort_values("Renewal_Rate"),

            x="Renewal_Rate",

            y= "New_IMD_Code",

            orientation="h",

            text="Renewal_Rate"

        )

        fig.update_traces(

            texttemplate="%{text:.1f}%"

        )

        fig.update_layout(

            height=420,

            xaxis_title="Renewal Rate (%)",

            yaxis_title=""

        )

        st.plotly_chart(

            fig,

            use_container_width=True

        )

    with right:

        st.subheader("🔴 Lowest Performing IMDs")

        bottom = (

            imd_summary

            .sort_values(

                "Renewal_Rate"

            )

            .head(10)

        )

        fig = px.bar(

            bottom,

            x="Renewal_Rate",

            y= "New_IMD_Code",

            orientation="h",

            text="Renewal_Rate"

        )

        fig.update_traces(

            texttemplate="%{text:.1f}%"

        )

        fig.update_layout(

            height=420,

            xaxis_title="Renewal Rate (%)",

            yaxis_title=""

        )

        st.plotly_chart(

            fig,

            use_container_width=True

        )

    st.divider()

    # ======================================================
    # Portfolio Analysis
    # ======================================================

    st.subheader("📈 Portfolio Size vs Renewal Rate")

    fig = px.scatter(

        imd_summary,

        x="Portfolio_Size",

        y="Renewal_Rate",

        size="Portfolio_Size",

        hover_name="New_IMD_Code",

        color="Renewal_Rate",

        height=500

    )

    st.plotly_chart(

        fig,

        use_container_width=True

    )

    st.divider()

    # ======================================================
    # Detailed IMD Table
    # ======================================================

    st.subheader("📋 IMD Performance Summary")

    display = (

        imd_summary

        .sort_values(

            "Renewal_Rate",

            ascending=False

        )

    )

    display["Renewal_Rate"] = (

        display["Renewal_Rate"]

        .round(1)

    )

    # display["Average_Premium"] = (

    #     display["Average_Premium"]

    #     .round(0)

    # )

    display["Average_Tenure"] = (

        display["Average_Tenure"]

        .round(1)

    )

    st.dataframe(

        display,

        use_container_width=True,

        hide_index=True

    )
    
    # ======================================================
    # LLM analysis
    # ======================================================

    imd_profiles = []

    for _, row in imd_summary.iterrows():
        
        imd_code = row["New_IMD_Code"]

        imd_df = channel_df[channel_df["New_IMD_Code"] == imd_code]

        remarks =  imd_df["REMARKS"].dropna().astype(str).tolist()

        remark_history = "\n".join(f"- {remark}" for remark in remarks)

        imd_profiles.append(
            
            f"""
        
        IMD Code: {imd_code}

        Portfolio Size: {int(row["Portfolio_Size"])}

        Renewal Rate: {row["Renewal_Rate"]}

        Average Vehicle Age: {row["Average_Vehicle_Age"]}

        Average Policy Tenure: {row["Average_Tenure"]}

        Remark History:

        {remark_history}

        ----------------------------------------
        """

        )

        prompt = f"""
 # ROLE

You are a Senior Motor Insurance Renewal Strategy Consultant preparing an executive operational audit for senior regional management.

The investigation workflow is controlled entirely by the application.

The Current Investigation is the ONLY source of truth.

Never add, infer or fabricate information beyond it.

Never output HTML.

---

# EXECUTION MODES

The application determines the execution mode.

## Investigation Mode

- Return ONLY the requested analytical tool call.
- Do not generate narrative text.
- Do not explain your actions.

## Report Generation Mode

- Generate the executive report using ONLY the Current Investigation.
- Do not call analytical tools.

---

# REPORT STYLE

Write like a McKinsey, Deloitte, EY or BCG consulting report.

The report should be:

- Executive-focused
- Concise
- Evidence-backed
- Business-oriented

Always explain the business implication rather than merely restating statistics.

---

# REPORT LENGTH

Target report length: **900–1200 words**.

Prioritize brevity over exhaustive explanation.

Keep paragraphs short.

Avoid repetition.

If statistics are already presented in a table, do not repeat the same numbers in the narrative.

---

# REPORT STRUCTURE

# Executive Summary

Maximum **6 one-sentence bullet points**.

Each bullet should summarise:

- Key Observation
- Supporting Evidence
- Business Impact

---

# Portfolio Assessment

Present ONE branch-level summary table.

| Metric | Value |

Include available branch statistics such as:

- Overall Renewal Rate
- Lost Percentage
- Follow-up Percentage
- Total IMDs Considered
- Average Portfolio Size
- Average Vehicle Age
- Average Policy Tenure

Follow the table with ONE short interpretation paragraph.

---

# IMD Performance Assessment

Group IMDs into:

## High Performing IMDs

## Average Performing IMDs

## Low Performing IMDs

For EACH category provide a table.

| IMD Code | Portfolio Size | Renewal Rate (%) | Vehicle Age Profile |

Include EVERY IMD contained in the Current Investigation.

After each table provide a maximum of THREE concise bullet points highlighting common characteristics.

---

# Lost Business Analysis

Present EVERY loss category whose percentage is greater than zero.

Use ONE consolidated table.

| Loss Category | Percentage | Representative Remarks | Business Interpretation | Controllability |

Where:

- Business Interpretation = one concise sentence.
- Controllability = Controllable / Partially Controllable / Unavoidable.

Representative Remarks must come directly from the Current Investigation.

---

# Strategic Recommendations

Present recommendations as a table.

| Stakeholder | Recommendation | Supporting Evidence | Expected Business Impact |

Stakeholders:

- Renewal Managers
- Relationship Managers
- Branch Operations
- Product Team
- Regional Management

Maximum TWO recommendations per stakeholder.

---

# WRITING RULES

- Every insight must be supported by evidence from the Current Investigation.
- Use tables for numerical information.
- Use bullet points for insights.
- Do not repeat observations.
- Do not repeat statistics already shown in tables.
- Focus on business impact and actionable insights.
- Never fabricate facts, statistics, IMD codes, historical remarks or competitor names.
- Output ONLY the final report in GitHub-Flavoured Markdown.

# MARKDOWN FORMAT

Use valid GitHub-Flavoured Markdown.

Correct heading syntax:

# Main Heading

## Section Heading

### Subsection Heading

Do NOT wrap headings in bold.

Correct:
## Executive Summary

Incorrect:
## **Executive Summary**

Incorrect:
# **Motor Insurance Renewal Strategy Audit**

Use bold only within normal paragraphs or table cells.

All tables must follow valid GitHub Markdown syntax.

Leave one blank line before and after every heading and table.
    """

    if "channel_reports" not in st.session_state:
        st.session_state.channel_reports = {}


    if st.button("🤖 AI Channel Analysis",key= f"ai_analysis_{channel}", use_container_width=True):
        with st.spinner("Analyzing channel performance..."):
            channel_analysis = asyncio.run(
                agent.generate_report(prompt, imd_channel = channel)
            )

            st.session_state.channel_reports[channel] = channel_analysis

    if channel in st.session_state.channel_reports:

        channel_analysis = st.session_state.channel_reports[channel]
        channel_analysis = clean_markdown_tables(channel_analysis)
        channel_analysis = re.sub(
            r'^(#{1,6})\s+\*\*(.*?)\*\*$',
            r'\1 \2',
            channel_analysis,
            flags=re.MULTILINE
        )
                
        print(repr(channel_analysis[:500]))

        html = f"""
        <style>

        body {{
            background-color: #0E1117;
            color: white;
            font-family: "Segoe UI", sans-serif;
            line-height: 1.7;
            padding: 20px;
                }}

        h1 {{
            color: #4FC3F7;
            border-bottom: 2px solid #444;
            padding-bottom: 8px;
        }}

        h2 {{
            color: #81C784;
            margin-top: 28px;
        }}

        h3 {{
            color: #BBDEFB;
        }}

        p, li {{
            color: white;
        }}

        strong {{
            color: white;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}

        th {{
            background: #1565C0;
            color: white;
            padding: 10px;
        }}

        td {{
            color: white;
            border: 1px solid #555;
            padding: 8px;
        }}

        tr:nth-child(even) {{
            background: #1E1E1E;
        }}

        code {{
            color: #FFD54F;
            background: #2D2D2D;
            padding: 2px 4px;
            border-radius: 4px;
        }}

        blockquote {{
            border-left: 4px solid #4FC3F7;
            padding-left: 12px;
            color: #E0E0E0;
        }}

        </style>

        {markdown.markdown(
            channel_analysis,
            extensions=["tables", "fenced_code", "nl2br"]
        )}
                """

        components.html(
            html,
            height = 1400,
            scrolling = True
        )

        print("CHANNEL ANALYSIS ACTUAL REPORT SENT BY LLM")
        print(channel_analysis)
            

    if channel in st.session_state.channel_reports:

        report_for_download = st.session_state.channel_reports[channel]

        word_file = markdown_to_word(
            report_for_download
        )

        st.download_button(
            "⬇ Download AI Report",
            data=word_file,
            file_name=f"{channel}_AI_Analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_ai_report_{channel}"
        )

        st.divider()

